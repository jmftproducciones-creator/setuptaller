# orden_docx.py
# Generador DOCX para órdenes - Setup (NBF Soft)
# Requisitos: pip install python-docx

import os
import re
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =========================
# Helpers de formato
# =========================

def _safe_str(x) -> str:
    if x is None:
        return ""
    return str(x).strip()

def _fmt_bool_si_no(v) -> str:
    return "SI" if str(v) in ("1", "True", "true", "SI", "Si", "sí", "Sí") else "NO"

def _fmt_fecha(v) -> str:
    """Acepta date/datetime/str. Devuelve 'YYYY-MM-DD' o ''."""
    if v is None:
        return ""
    if isinstance(v, datetime):
        return v.date().isoformat()
    if isinstance(v, date):
        return v.isoformat()
    s = _safe_str(v)
    # si viene "2025-12-18 00:00:00" -> recorto
    if len(s) >= 10 and re.match(r"^\d{4}-\d{2}-\d{2}", s):
        return s[:10]
    return s

def _fmt_hora(v) -> str:
    """Acepta time/datetime/str. Devuelve 'HH:MM' o ''."""
    if v is None:
        return ""
    if isinstance(v, datetime):
        return v.strftime("%H:%M")
    s = _safe_str(v)
    if re.match(r"^\d{2}:\d{2}", s):
        return s[:5]
    return s

def _fmt_fecha_hora(fecha, hora) -> str:
    f = _fmt_fecha(fecha)
    h = _fmt_hora(hora)
    if f and h:
        return f"{f} {h}"
    if f:
        return f
    if h:
        return h
    return ""

def _sanitize_filename(name: str) -> str:
    name = _safe_str(name)
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name[:140] if len(name) > 140 else name


# =========================
# Construcción de DOCX
# =========================

def _set_default_font(doc: Document, font_name="Calibri", font_size=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(font_size)

def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_kv_table(doc: Document, rows_2col):
    """
    rows_2col: list of tuples (k, v)
    """
    table = doc.add_table(rows=len(rows_2col), cols=2)
    table.style = "Table Grid"
    # anchos
    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(12.0)

    for i, (k, v) in enumerate(rows_2col):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)

        c0.paragraphs[0].add_run(_safe_str(k)).bold = True
        c1.paragraphs[0].add_run(_safe_str(v))

    doc.add_paragraph()  # espacio

def _add_section(doc: Document, title: str, content: str):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph(_safe_str(content) or "—")
    doc.add_paragraph()  # espacio

def _join_multiline(value) -> str:
    """
    Si viene como lista, lo une con saltos.
    Si viene string con separadores (|, \n), lo deja prolijo.
    """
    if value is None:
        return ""
    if isinstance(value, (list, tuple)):
        return "\n".join([_safe_str(x) for x in value if _safe_str(x)])
    s = _safe_str(value)
    # normalizar separadores comunes
    s = s.replace(" | ", "\n").replace("|", "\n")
    # colapsar saltos extra
    s = re.sub(r"\n{3,}", "\n\n", s).strip()
    return s


# =========================
# API principal
# =========================

def generar_docx_orden(orden: dict, output_dir: str = "docx_ordenes") -> str:
    """
    Genera un DOCX para una orden y devuelve la ruta absoluta del archivo generado.
    Espera un dict con claves típicas de tu SELECT (o.* + joins):
      - id, fecha, hora_ingreso, estado, importe, ...
      - nombre_contacto, telefono_contacto (opcional)
      - equipo_texto, serie_texto (opcional)
      - presupuesto_aprobado
      - fecha_salida/hora_salida, fecha_regreso/hora_regreso, fecha_retiro/hora_retiro
      - falla, reparacion, repuestos, observaciones, accesorios
    """

    os.makedirs(output_dir, exist_ok=True)

    oid = orden.get("id") or orden.get("nro") or ""
    estado = _safe_str(orden.get("estado"))
    cliente = _safe_str(orden.get("nombre_contacto") or orden.get("cliente") or "")
    tel = _safe_str(orden.get("telefono_contacto") or orden.get("telefono") or "")
    equipo = _safe_str(orden.get("equipo_texto") or orden.get("equipo") or "")
    serie = _safe_str(orden.get("serie_texto") or orden.get("serie") or "")
    importe = orden.get("importe")
    try:
        importe_txt = f"{float(str(importe).replace(',', '.')):.2f}"
    except:
        importe_txt = _safe_str(importe)

    pa_txt = _fmt_bool_si_no(orden.get("presupuesto_aprobado"))

    ingreso_txt = _fmt_fecha_hora(orden.get("fecha"), orden.get("hora_ingreso"))
    salida_txt  = _fmt_fecha_hora(orden.get("fecha_salida"), orden.get("hora_salida"))
    regreso_txt = _fmt_fecha_hora(orden.get("fecha_regreso"), orden.get("hora_regreso"))
    retiro_txt  = _fmt_fecha_hora(orden.get("fecha_retiro"), orden.get("hora_retiro"))

    falla_txt = _join_multiline(orden.get("falla"))
    reparacion_txt = _join_multiline(orden.get("reparacion"))
    repuestos_txt = _join_multiline(orden.get("repuestos"))
    observaciones_txt = _safe_str(orden.get("observaciones"))
    accesorios_txt = _safe_str(orden.get("accesorios"))

    doc = Document()
    _set_default_font(doc, "Calibri", 11)

    # Márgenes
    sec = doc.sections[0]
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    # Títulos
    _add_title(doc, "SETUP - ORDEN DE SERVICIO")
    _add_subtitle(doc, f"Orden N° {oid}")

    # Datos principales
    rows = [
        ("Estado", estado),
        ("Presupuesto aprobado", pa_txt),
        ("Ingreso", ingreso_txt),
        ("Salida", salida_txt or "—"),
        ("Regreso", regreso_txt or "—"),
        ("Retiro", retiro_txt or "—"),
        ("Cliente / Contacto", cliente or "—"),
        ("Teléfono", tel or "—"),
        ("Equipo", equipo or "—"),
        ("Serie / S/N", serie or "—"),
        ("Importe", importe_txt or "0.00"),
    ]
    _add_kv_table(doc, rows)

    # Secciones de texto
    _add_section(doc, "Falla", falla_txt)
    _add_section(doc, "Reparación", reparacion_txt)
    _add_section(doc, "Repuestos", repuestos_txt)
    _add_section(doc, "Accesorios", accesorios_txt)
    _add_section(doc, "Observaciones", observaciones_txt)

    # Pie
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NBF Soft • Setup")
    run.italic = True
    run.font.size = Pt(9)

    # Nombre del archivo
    cliente_fname = _sanitize_filename(cliente) or "cliente"
    filename = _sanitize_filename(f"orden_{oid}_{cliente_fname}.docx")
    fullpath = os.path.abspath(os.path.join(output_dir, filename))

    doc.save(fullpath)
    return fullpath
